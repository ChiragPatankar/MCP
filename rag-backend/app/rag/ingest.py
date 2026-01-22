"""
Document ingestion and parsing pipeline.
Supports PDF, DOCX, TXT, and Markdown files.
"""
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import markdown
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import chardet
import re
from dataclasses import dataclass
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """Represents a parsed document with text and metadata."""
    text: str
    file_name: str
    file_type: str
    page_count: Optional[int] = None
    page_map: Optional[Dict[int, int]] = None  # char_position -> page_number
    metadata: Dict[str, Any] = None


class DocumentParser:
    """
    Parses various document formats into plain text.
    Preserves page information for citations.
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.md', '.markdown'}
    
    def __init__(self):
        self.parsers = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx,  # Try docx parser for doc files
            '.txt': self._parse_text,
            '.md': self._parse_markdown,
            '.markdown': self._parse_markdown,
        }
    
    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse a document file into text.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ParsedDocument with extracted text and metadata
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")
        
        parser = self.parsers.get(ext)
        if not parser:
            raise ValueError(f"No parser available for: {ext}")
        
        logger.info(f"Parsing document: {file_path.name} ({ext})")
        return parser(file_path)
    
    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """Parse PDF file with page tracking."""
        try:
            doc = fitz.open(file_path)
            text_parts = []
            page_map = {}
            current_pos = 0
            
            for page_num, page in enumerate(doc, start=1):
                page_text = page.get_text("text")
                if page_text.strip():
                    # Record where this page starts
                    page_map[current_pos] = page_num
                    text_parts.append(page_text)
                    current_pos += len(page_text) + 2  # +2 for separator
            
            page_count = len(doc)
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            
            return ParsedDocument(
                text=self._clean_text(full_text),
                file_name=file_path.name,
                file_type="pdf",
                page_count=page_count,
                page_map=page_map,
                metadata={"source": str(file_path)}
            )
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise
    
    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """Parse DOCX file."""
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            full_text = "\n\n".join(paragraphs)
            
            return ParsedDocument(
                text=self._clean_text(full_text),
                file_name=file_path.name,
                file_type="docx",
                metadata={"source": str(file_path)}
            )
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise
    
    def _parse_text(self, file_path: Path) -> ParsedDocument:
        """Parse plain text file with encoding detection."""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                detected = chardet.detect(raw_data)
                encoding = detected.get('encoding', 'utf-8')
            
            # Read with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                text = f.read()
            
            return ParsedDocument(
                text=self._clean_text(text),
                file_name=file_path.name,
                file_type="txt",
                metadata={"source": str(file_path), "encoding": encoding}
            )
        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {e}")
            raise
    
    def _parse_markdown(self, file_path: Path) -> ParsedDocument:
        """Parse Markdown file, converting to plain text."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                md_content = f.read()
            
            # Convert markdown to HTML, then strip tags
            html = markdown.markdown(md_content)
            text = self._strip_html_tags(html)
            
            # Also keep the original markdown structure for better context
            # Remove markdown syntax but keep structure
            clean_md = self._clean_markdown(md_content)
            
            return ParsedDocument(
                text=self._clean_text(clean_md),
                file_name=file_path.name,
                file_type="markdown",
                metadata={"source": str(file_path)}
            )
        except Exception as e:
            logger.error(f"Error parsing Markdown {file_path}: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Replace multiple whitespace with single space
        text = re.sub(r'[ \t]+', ' ', text)
        # Replace multiple newlines with double newline
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Remove leading/trailing whitespace from lines
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        # Remove leading/trailing whitespace
        return text.strip()
    
    def _strip_html_tags(self, html: str) -> str:
        """Remove HTML tags from text."""
        clean = re.sub(r'<[^>]+>', '', html)
        return clean
    
    def _clean_markdown(self, md_text: str) -> str:
        """Clean markdown syntax while preserving structure."""
        # Remove code blocks but keep content
        md_text = re.sub(r'```[\s\S]*?```', '', md_text)
        md_text = re.sub(r'`([^`]+)`', r'\1', md_text)
        
        # Convert headers to plain text with emphasis
        md_text = re.sub(r'^#{1,6}\s+(.+)$', r'\1:', md_text, flags=re.MULTILINE)
        
        # Remove bold/italic markers
        md_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', md_text)
        md_text = re.sub(r'\*([^*]+)\*', r'\1', md_text)
        md_text = re.sub(r'__([^_]+)__', r'\1', md_text)
        md_text = re.sub(r'_([^_]+)_', r'\1', md_text)
        
        # Remove links but keep text
        md_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', md_text)
        
        # Remove images
        md_text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', '', md_text)
        
        # Convert lists to plain text
        md_text = re.sub(r'^[\*\-\+]\s+', '• ', md_text, flags=re.MULTILINE)
        md_text = re.sub(r'^\d+\.\s+', '', md_text, flags=re.MULTILINE)
        
        return md_text


# Global parser instance
parser = DocumentParser()

